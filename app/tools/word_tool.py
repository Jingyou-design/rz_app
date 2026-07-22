from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

import os

def concat_py_files(files, encoding='utf-8', order=None):
    """
    files: 文件路径列表，或文件夹路径(str)
    order: 可选，字典，格式如 {"files": ["main.py", "config.py"]}。
           提供时按 files 列表顺序拼接，未列出的文件按字母序追加末尾，
           列表中不存在于 files 的文件名跳过并警告。
    """
    if isinstance(files, str) and os.path.isdir(files):
        files = [
            os.path.join(files, f)
            for f in sorted(os.listdir(files))
            if f.endswith('.py')
        ]

    if order and isinstance(order, dict):
        order_list = order.get('files', [])
        basename_map = {os.path.basename(f): f for f in files}
        ordered = []
        used = set()
        for name in order_list:
            if name in basename_map:
                ordered.append(basename_map[name])
                used.add(name)
            else:
                print(f'跳过（order 中但不存在）: {name}')
        remaining = sorted(
            f for f in files if os.path.basename(f) not in used
        )
        files = ordered + remaining

    parts = []
    for path in files:
        if not os.path.isfile(path):
            print(f'跳过（不存在）: {path}')
            continue
        with open(path, 'r', encoding=encoding) as f:
            content = f.read()
        parts.append(content)
        if not content.endswith('\n'):
            parts.append('')

    return '\n'.join(parts)

def fill_new_content(docx_path, output_path, text):

    doc = Document(docx_path)
    body = doc.element.body

    # 保留 sectPr，删除其余内容
    sectPr = body.find(qn('w:sectPr'))
    sectPr_copy = deepcopy(sectPr) if sectPr is not None else None
    for child in list(body):
        body.remove(child)

    # 每行一个独立段落，合并连续空行为单个（匹配参考文档格式）
    lines = text.split('\n')
    merged = []
    prev_empty = False
    for line in lines:
        if line == '':
            if prev_empty:
                continue
            prev_empty = True
        else:
            prev_empty = False
        merged.append(line)
    for line in merged:
        p = doc.add_paragraph()
        # 固定行距（匹配参考文档 line=280 lineRule=exact）
        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '280')
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)
        # run 不设字体，继承 Normal 样式
        p.add_run(line)

    # sectPr 放回 body 末尾
    if sectPr_copy is not None:
        body.append(sectPr_copy)

    doc.save(output_path)
